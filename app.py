from llama_index.core import PromptTemplate
from llama_index.llms.groq import Groq
from dotenv import load_dotenv
import os
import streamlit as st
from docx import Document


# Load environment variables
load_dotenv()

# Load Google API key from environment
GROOQ_API_KEY = os.getenv("GROOQ_API_KEY")

# Initialize Gemini
llm = Groq(model="llama3-70b-8192", api_key=GROOQ_API_KEY)

# Streamlit app
st.set_page_config(layout="wide", page_title="HR Assistant")
st.title("Kucing Belang")

# Sidebar
with st.sidebar:
    st.sidebar.header("Input your HR Data")

    company_name = st.text_input(
        label="Company Name",
        help="Input your company name",
    )
    company_data = st.text_area(
        label="Masukan data perusahaan Anda",
        height=100,
        placeholder="Tuliskan profile perusahaan Anda.",
        help="Data ini anda akan dipergunakan untuk mendapatkan konteks dalam pembuatan rekomendasi yang kami berikan"
    )

    position_data = st.text_input(
        label="Nama Jabatan",
        help="Input nama jabatan yang dicari",
        placeholder="contoh:Frontend Developer"
    )

    job_description_data = st.text_area(
        label="Deskripsi Jabatan",
        help="Input deskripsi jabatan yang dicari",
        placeholder="Tuliskan deskripsi jabatan yang dicari"
    )

    contact_us = st.text_area(
        label="Contact Us",
        help="Cara menghubungi pesrushaan",
        placeholder="Email perusahaan,nomer telfon atau alamat telfon"
    )
    analyze_button = st.button("Start Create Iklan")
    

tab1, tab2, tab3, tab4 = st.tabs(["Iklan", "Job Test", "Pertanyaan Interview", "Test Kompetensi"])

def generate_content(prompt):
    response = llm.complete(prompt)
    return response

def save_to_word(content, filename="HR_Assistant_Output.docx"):
    doc = Document()
    doc.add_heading('HR Assistant Output', 0)
    doc.add_paragraph(content)
    doc.save(filename)

def load_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        competencies_data = file.read()
    return competencies_data

if analyze_button: 
        
    with tab1:
        # Generate Iklan
        with st.spinner("Generating content..."):  
            st.subheader(f"Job Ad for {position_data}")
            ad_prompt=f"Buatkan iklan untuk posisi untuk {company_name} {position_data} di {company_data} dengan deskripsi pekerjaan berikut: {job_description_data},Buat dalam bahasa indonesia dengan struktur pembukaan, job deskripsi, job requirement, dan cara melamar dengan menhubungi {contact_us}"
            Job_advertising = generate_content(ad_prompt)
            st.markdown(f'<div class="stock-analysis">{Job_advertising}</div>', unsafe_allow_html=True)
    with tab2:
        # Generate Job test
        with st.spinner("Generating content..."):  
            st.subheader(f"Job Testing for {position_data}")
            jobtest_prompt=f"Buatkan job test yang terdiri atas maksimal 4 pertanyaan untuk posisi untuk {company_name} {position_data} di {company_data} dengan deskripsi pekerjaan berikut: {job_description_data}, pastikan juga anda memasukan data {Job_advertising} sebagai pertimabangan dalam pembuatan pertanyaan. Buat dalam bahasa indonesia dengan struktur:instruksi, cara menjawab dengan mengirimkan email ke {contact_us}."
            Job_test = generate_content(jobtest_prompt)
            st.markdown(f'<div class="stock-analysis">{Job_test}</div>', unsafe_allow_html=True)

            st.subheader("Keypoint Anwer")
            answer_prompt=f"berdasarkan pertanyaan dari {Job_test}, buatkan 3 keyword yang harus ada dalam masing masing jawaban."
            test_answer = generate_content(answer_prompt)
            st.markdown(f'<div class="stock-analysis">{test_answer}</div>', unsafe_allow_html=True)

    with tab3:
        # Generate interview
        with st.spinner("Generating content..."):  
            st.subheader(f"Interview qusetion for {position_data}")
            interview_prompt=f"berdasarkan pertanyaan dari {Job_test} dan {job_description_data} dan {Job_test}, buatkan 3 pertanyaan interview yang harus ada dalam masing masing jawaban.buat dalam bahasa indoensia yang baik dan benar"
            interview = generate_content(interview_prompt)
            st.markdown(f'<div class="stock-analysis">{interview}</div>', unsafe_allow_html=True)

    with tab4:
        with st.spinner("Generating content..."):
            competencies_book = load_text_file('competencies.txt')
            competencies_prompt=(f'berdasarkan {competencies_book}, pilihlah dan jelaskan kompetensi yang harus dimiliki oleh {position_data}. Hanya gunakan data dari dokumen yang disertakan, jangan menggunakan data lainnya, tuliskan dalam bahassa indonesia  yang jelas.')
            competency_content = generate_content(competencies_prompt)
            st.subheader(f"Competency Test for {position_data}")
            st.markdown(f'<div class="stock-analysis">{competency_content}</div>', unsafe_allow_html=True)

            